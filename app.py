import io
import json
from typing import Generator, Iterable, List, Optional, Tuple

import streamlit as st

# Optional imports guarded to improve UX if modules are missing
try:
	# Local LLM via Ollama
	import ollama
except Exception as e:  # pragma: no cover
	ollama = None

try:
	from langdetect import detect, DetectorFactory
	DetectorFactory.seed = 42
except Exception:  # pragma: no cover
	detect = None

try:
	import chardet
except Exception:  # pragma: no cover
	chardet = None

try:
	from docx import Document as DocxDocument
except Exception:  # pragma: no cover
	DocxDocument = None

try:
	from pypdf import PdfReader
except Exception:  # pragma: no cover
	PdfReader = None


# Common ISO 639-1 languages map
LANGUAGES: List[Tuple[str, str]] = [
	("auto", "Auto-detect"),
	("af", "Afrikaans"),
	("ar", "Arabic"),
	("bn", "Bengali"),
	("bg", "Bulgarian"),
	("ca", "Catalan"),
	("zh-cn", "Chinese (Simplified)"),
	("zh-tw", "Chinese (Traditional)"),
	("hr", "Croatian"),
	("cs", "Czech"),
	("da", "Danish"),
	("nl", "Dutch"),
	("en", "English"),
	("et", "Estonian"),
	("fi", "Finnish"),
	("fr", "French"),
	("de", "German"),
	("el", "Greek"),
	("he", "Hebrew"),
	("hi", "Hindi"),
	("hu", "Hungarian"),
	("id", "Indonesian"),
	("it", "Italian"),
	("ja", "Japanese"),
	("ko", "Korean"),
	("lv", "Latvian"),
	("lt", "Lithuanian"),
	("ms", "Malay"),
	("no", "Norwegian"),
	("fa", "Persian"),
	("pl", "Polish"),
	("pt", "Portuguese"),
	("ro", "Romanian"),
	("ru", "Russian"),
	("sr", "Serbian"),
	("sk", "Slovak"),
	("sl", "Slovenian"),
	("es", "Spanish"),
	("sv", "Swedish"),
	("ta", "Tamil"),
	("te", "Telugu"),
	("th", "Thai"),
	("tr", "Turkish"),
	("uk", "Ukrainian"),
	("ur", "Urdu"),
	("vi", "Vietnamese"),
]

LANG_CODE_TO_NAME = {code: name for code, name in LANGUAGES}
LANG_NAME_TO_CODE = {name: code for code, name in LANGUAGES}


def language_name_from_code(code: str) -> str:
	code = (code or "").lower()
	return LANG_CODE_TO_NAME.get(code, code)


def installed_ollama_models() -> List[str]:
	if ollama is None:
		return []
	try:
		resp = ollama.list()
		models = [m.get("name") for m in resp.get("models", []) if m.get("name")]
		return sorted(models)
	except Exception:
		return []

def is_translation_model_name(model_name: str) -> bool:
	name = (model_name or "").lower()
	# Exclude obvious non-chat / non-translation models
	block_keywords = [
		"embed", "nomic", "llava", "vision", "clip", "whisper", "asr", "tts",
		"guard", "llama-guard", "classify", "moderate", "sd", "image",
		"coder", "codeqwen", "deepseek-coder", "codellama", "starcoder",
	]
	for kw in block_keywords:
		if kw in name:
			return False
	# Include common multilingual/chat instruct families
	allow_keywords = [
		"qwen", "llama", "mistral", "mixtral", "gemma", "yi", "zephyr",
		"vicuna", "openchat", "solar", "phi", "glm", "bloomz", "xglm",
	]
	return any(kw in name for kw in allow_keywords)

def filter_translation_models(models: List[str]) -> List[str]:
	return [m for m in models if is_translation_model_name(m)]


def auto_detect_language(text: str) -> Optional[str]:
	if not text or detect is None:
		return None
	try:
		code = detect(text)
		# Map general 'zh' to a default variant
		if code == "zh":
			return "zh-cn"
		return code
	except Exception:
		return None


def build_translation_system_prompt() -> str:
	return (
		"You are a professional translator. Translate the user's input exactly and faithfully.\n"
		"- Preserve original meaning, tone, and formatting (line breaks, lists, punctuation).\n"
		"- Do not add explanations or notes. Output only the translated text.\n"
		"- Keep numbers, URLs, and code blocks unchanged. Maintain inline markup if present.\n"
	)


def build_translation_user_prompt(text: str, src_lang: str, tgt_lang: str) -> str:
	return (
		f"Source language: {language_name_from_code(src_lang)}\n"
		f"Target language: {language_name_from_code(tgt_lang)}\n\n"
		f"Translate the following content. Output only the translated text:\n\n{text}"
	)


def chat_complete_stream(
	model: str,
	system_prompt: str,
	user_prompt: str,
	temperature: float = 0.2,
) -> Generator[str, None, None]:
	for part in ollama.chat(
		model=model,
		messages=[
			{"role": "system", "content": system_prompt},
			{"role": "user", "content": user_prompt},
		],
		options={"temperature": temperature},
		stream=True,
	):
		# Newer ollama returns chunks with message.content
		try:
			chunk = part.get("message", {}).get("content", "")
			if chunk:
				yield chunk
		except Exception:
			# Fallback for legacy field 'response'
			chunk = part.get("response", "")
			if chunk:
				yield chunk


def chat_complete(
	model: str,
	system_prompt: str,
	user_prompt: str,
	temperature: float = 0.2,
) -> str:
	resp = ollama.chat(
		model=model,
		messages=[
			{"role": "system", "content": system_prompt},
			{"role": "user", "content": user_prompt},
		],
		options={"temperature": temperature},
	)
	# Newer ollama returns message.content
	message = resp.get("message", {})
	content = message.get("content")
	if content:
		return content
	# Fallback for legacy field 'response'
	return resp.get("response", "")


def translate_text(
	text: str,
	model: str,
	src_lang: str,
	tgt_lang: str,
	temperature: float = 0.2,
	stream: bool = True,
) -> Iterable[str] | str:
	if not text.strip():
		return ""
	if src_lang == "auto":
		detected = auto_detect_language(text) or "en"
		src_lang = detected
	if src_lang == tgt_lang:
		return text
	system_prompt = build_translation_system_prompt()
	user_prompt = build_translation_user_prompt(text, src_lang, tgt_lang)
	if stream:
		return chat_complete_stream(model, system_prompt, user_prompt, temperature)
	return chat_complete(model, system_prompt, user_prompt, temperature)


def chunk_paragraphs(paragraphs: List[str], max_chars: int = 1800) -> List[str]:
	chunks: List[str] = []
	current: List[str] = []
	current_len = 0
	for p in paragraphs:
		if not p:
			# Preserve empty paragraph boundaries
			if current_len + 1 > max_chars and current:
				chunks.append("\n".join(current))
				current = []
				current_len = 0
			current.append("")
			current_len += 1
			continue
		if current_len + len(p) + 1 > max_chars and current:
			chunks.append("\n".join(current))
			current = [p]
			current_len = len(p)
		else:
			current.append(p)
			current_len += len(p) + 1
	if current:
		chunks.append("\n".join(current))
	return chunks


def read_text_file(uploaded) -> str:
	data = uploaded.read()
	if not data:
		return ""
	if chardet:
		enc = chardet.detect(data).get("encoding") or "utf-8"
	else:
		enc = "utf-8"
	try:
		return data.decode(enc, errors="replace")
	except Exception:
		return data.decode("utf-8", errors="replace")


def extract_docx_text(uploaded) -> List[str]:
	if DocxDocument is None:
		raise RuntimeError("python-docx is not installed")
	doc = DocxDocument(uploaded)
	return [p.text for p in doc.paragraphs]


def extract_pdf_text(uploaded) -> List[str]:
	if PdfReader is None:
		raise RuntimeError("pypdf is not installed")
	reader = PdfReader(uploaded)
	pages_text: List[str] = []
	for page in reader.pages:
		pages_text.append(page.extract_text() or "")
	return pages_text


def write_docx(paragraphs: List[str]) -> bytes:
	if DocxDocument is None:
		raise RuntimeError("python-docx is not installed")
	doc = DocxDocument()
	for p in paragraphs:
		doc.add_paragraph(p)
	buf = io.BytesIO()
	doc.save(buf)
	return buf.getvalue()


def translate_large_text_in_chunks(
	text: str,
	model: str,
	src_lang: str,
	tgt_lang: str,
	temperature: float,
) -> str:
	paras = text.splitlines()
	chunks = chunk_paragraphs(paras, max_chars=1800)
	out_parts: List[str] = []
	for chunk in chunks:
		translated = translate_text(
			chunk, model=model, src_lang=src_lang, tgt_lang=tgt_lang, temperature=temperature, stream=False
		)
		out_parts.append(str(translated))
	return "\n".join(out_parts)


def page_sidebar() -> Tuple[str, float, bool]:
	st.sidebar.header("Model & Settings")
	models = installed_ollama_models()
	show_only_trans = st.sidebar.toggle("Show translation-capable only", value=True, key="only_trans_models")
	if show_only_trans:
		models = filter_translation_models(models)
	default_suggestions = [
		"qwen2.5:7b-instruct",
		"qwen2.5:14b-instruct",
		"llama3.1:8b",
		"qwen2.5:3b-instruct",
		"mistral:7b-instruct",
		"gemma2:9b-instruct",
	]
	model_options = models if models else default_suggestions
	model = st.sidebar.selectbox("Ollama model", model_options, index=0)
	temperature = st.sidebar.slider("Creativity (temperature)", 0.0, 1.0, 0.2, 0.05)
	stream = st.sidebar.toggle("Stream output", value=True)
	if not models:
		st.sidebar.info(
			"No local models found. Use `ollama pull qwen2.5:7b-instruct` or another model, "
			"then restart this app."
		)
	return model, temperature, stream


def language_selectors(detected_code: Optional[str], key_ns: str = "") -> Tuple[str, str]:
	col1, col2 = st.columns(2)
	with col1:
		src = st.selectbox(
			"Source language",
			[name for _, name in LANGUAGES],
			index=0,
			key=f"{key_ns}src_lang",
		)
	with col2:
		target_names = [LANG_CODE_TO_NAME[c] for c in LANG_CODE_TO_NAME if c != "auto"]
		default_target_index = target_names.index("English") if "English" in target_names else 0
		tgt = st.selectbox(
			"Target language",
			target_names,
			index=default_target_index,
			key=f"{key_ns}tgt_lang",
		)
	src_code = LANG_NAME_TO_CODE.get(src, "auto")
	tgt_code = next((code for code, name in LANGUAGES if name == tgt), "en")
	if detected_code and src_code == "auto":
		st.caption(f"Detected source: {language_name_from_code(detected_code)}")
	return src_code, tgt_code


def ui_text_translation(model: str, temperature: float, stream: bool):
	st.subheader("Text Translation")
	text = st.text_area("Enter text", height=180, placeholder="Paste or type your text to translate…")
	detected = auto_detect_language(text) if text else None
	src_code, tgt_code = language_selectors(detected, key_ns="text_")
	translate_clicked = st.button("Translate")
	if translate_clicked:
		if not text.strip():
			st.warning("Please enter some text to translate.")
			return
		if ollama is None:
			st.error("The 'ollama' Python package is not available. Install dependencies and retry.")
			return
		st.write(f"Translating {language_name_from_code(src_code)} → {language_name_from_code(tgt_code)} using `{model}`")
		output_placeholder = st.empty()
		if stream:
			chunks = translate_text(
				text,
				model=model,
				src_lang=src_code,
				tgt_lang=tgt_code,
				temperature=temperature,
				stream=True,
			)
			collected = []
			for ch in chunks:  # type: ignore
				collected.append(ch)
				output_placeholder.markdown("".join(collected))
		else:
			result = translate_text(
				text,
				model=model,
				src_lang=src_code,
				tgt_lang=tgt_code,
				temperature=temperature,
				stream=False,
			)
			output_placeholder.markdown(str(result))


def ui_file_translation(model: str, temperature: float):
	st.subheader("File Translation")
	uploaded = st.file_uploader("Upload a file (.txt, .md, .docx, .pdf)", type=["txt", "md", "docx", "pdf"])
	output_format = st.selectbox("Output format", ["docx", "txt"], index=0)
	if uploaded is None:
		return
	filename = uploaded.name
	ext = (filename.split(".")[-1] or "").lower()

	# Choose target language, source optional auto
	src_code, tgt_code = language_selectors(None, key_ns="file_")

	if st.button("Translate File"):
		if ollama is None:
			st.error("The 'ollama' Python package is not available. Install dependencies and retry.")
			return
		try:
			if ext in {"txt", "md"}:
				text = read_text_file(uploaded)
				detected = auto_detect_language(text) or "en"
				if src_code == "auto":
					src_code_use = detected
				else:
					src_code_use = src_code
				result = translate_large_text_in_chunks(
					text=text,
					model=model,
					src_lang=src_code_use,
					tgt_lang=tgt_code,
					temperature=temperature,
				)
				# Preview
				st.markdown("Preview")
				st.text_area("Translated preview", result, height=300, key="file_preview_text")
				# Downloads (both formats)
				out_base = filename.rsplit(".", 1)[0]
				data_txt = result.encode("utf-8")
				data_docx = write_docx(result.splitlines())
				col_dl1, col_dl2 = st.columns(2)
				with col_dl1:
					st.download_button(
						"Download .txt",
						data_txt,
						file_name=out_base + f".{tgt_code}.txt",
						mime="text/plain",
					)
				with col_dl2:
					st.download_button(
						"Download .docx",
						data_docx,
						file_name=out_base + f".{tgt_code}.docx",
						mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
					)

			elif ext == "docx":
				paras = extract_docx_text(uploaded)
				# Detect source based on joined content if auto
				joined = "\n".join(paras)
				detected = auto_detect_language(joined) or "en"
				src_code_use = detected if src_code == "auto" else src_code
				chunks = chunk_paragraphs(paras, max_chars=1800)
				out_paras: List[str] = []
				for chunk in chunks:
					translated = translate_text(
						chunk,
						model=model,
						src_lang=src_code_use,
						tgt_lang=tgt_code,
						temperature=temperature,
						stream=False,
					)
					out_paras.extend(str(translated).splitlines())
				result_str = "\n".join(out_paras)
				# Preview
				st.markdown("Preview")
				st.text_area("Translated preview", result_str, height=300, key="file_preview_docx")
				# Downloads (both formats)
				out_base = filename.rsplit(".", 1)[0]
				data_txt = result_str.encode("utf-8")
				data_docx = write_docx(out_paras)
				col_dl1, col_dl2 = st.columns(2)
				with col_dl1:
					st.download_button(
						"Download .txt",
						data_txt,
						file_name=out_base + f".{tgt_code}.txt",
						mime="text/plain",
					)
				with col_dl2:
					st.download_button(
						"Download .docx",
						data_docx,
						file_name=out_base + f".{tgt_code}.docx",
						mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
					)

			elif ext == "pdf":
				pages = extract_pdf_text(uploaded)
				joined = "\n\n".join(pages)
				detected = auto_detect_language(joined) or "en"
				src_code_use = detected if src_code == "auto" else src_code
				result = translate_large_text_in_chunks(
					text=joined,
					model=model,
					src_lang=src_code_use,
					tgt_lang=tgt_code,
					temperature=temperature,
				)
				# Preview
				st.markdown("Preview")
				st.text_area("Translated preview", result, height=300, key="file_preview_pdf")
				# Downloads (both formats)
				out_base = filename.rsplit(".", 1)[0]
				data_txt = result.encode("utf-8")
				data_docx = write_docx(result.splitlines())
				col_dl1, col_dl2 = st.columns(2)
				with col_dl1:
					st.download_button(
						"Download .txt",
						data_txt,
						file_name=out_base + f".{tgt_code}.txt",
						mime="text/plain",
					)
				with col_dl2:
					st.download_button(
						"Download .docx",
						data_docx,
						file_name=out_base + f".{tgt_code}.docx",
						mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
					)

			else:
				st.error("Unsupported file type.")
		except Exception as e:
			st.error(f"Translation failed: {e}")


def human_readable_bytes(num: int) -> str:
	for unit in ["B", "KB", "MB", "GB", "TB"]:
		if num < 1024:
			return f"{num:.1f} {unit}"
		num /= 1024
	return f"{num:.1f} PB"


def installed_models_detailed() -> List[dict]:
	if ollama is None:
		return []
	try:
		return ollama.list().get("models", [])
	except Exception:
		return []

def quick_translation_probe(model: str) -> Tuple[bool, str]:
	try:
		resp = chat_complete(
			model=model,
			system_prompt=build_translation_system_prompt(),
			user_prompt="Translate the following to English. Output only the translation.\n\nBonjour",
			temperature=0.0,
		).strip()
		return ("hello" in resp.lower(), resp)
	except Exception as e:
		return (False, f"error: {e}")


def ui_models_tab():
	st.subheader("Models")
	col_left, col_right = st.columns([2, 1])
	with col_left:
		st.markdown("Installed models")
		models = installed_models_detailed()
		model_names = [m.get("name", "") for m in models]
		trans_names = set(filter_translation_models(model_names))
		if not models:
			st.info("No local models found.")
		else:
			for m in models:
				name = m.get("name", "")
				size = m.get("size", 0)
				digest = m.get("digest", "")[:12]
				tag = "✅ trans" if name in trans_names else "—"
				st.markdown(f"- `{name}` — {human_readable_bytes(size)} (id: `{digest}`) {tag}")
	with col_right:
		st.markdown("Find more models")
		st.markdown("[Browse Ollama Library](https://ollama.com/models)")
		st.markdown("[Hugging Face (Ollama)](https://huggingface.co/models?library=ollama)")
		st.caption("Open links to explore and copy model names.")

	st.divider()
	st.markdown("Pull a model by name")
	model_to_pull = st.text_input("Model name (e.g., qwen2.5:7b-instruct)")
	recs = [
		"qwen2.5:7b-instruct",
		"llama3.1:8b",
		"mistral:7b-instruct",
		"gemma2:9b-instruct",
	]
	st.caption("Quick picks:")
	rec_cols = st.columns(len(recs))
	for i, rec in enumerate(recs):
		if rec_cols[i].button(f"Pull {rec}"):
			model_to_pull = rec
	if st.button("Pull model") and model_to_pull:
		if ollama is None:
			st.error("Ollama client not available.")
		else:
			with st.spinner(f"Pulling {model_to_pull}…"):
				try:
					for _status in ollama.pull(model_to_pull, stream=True):  # type: ignore
						pass
					st.success(f"Pulled {model_to_pull}. Refresh the page if it doesn't appear.")
				except Exception as e:
					st.error(f"Failed to pull: {e}")

	st.divider()
	st.markdown("Validate translation capability (quick probe)")
	all_installed = installed_ollama_models()
	if not all_installed:
		st.caption("No models available to validate.")
		return
	model_to_test = st.selectbox("Select a model to test", all_installed, key="probe_model")
	if st.button("Run probe"):
		ok, resp = quick_translation_probe(model_to_test)
		if ok:
			st.success(f"Probe passed. Output: “{resp}”")
		else:
			st.warning(f"Probe did not pass. Output: “{resp}”")


def ui_compare_models(temperature: float):
	st.subheader("Compare Models")
	models = filter_translation_models(installed_ollama_models())
	if len(models) < 2:
		st.info("Install at least two models to compare. Use the Models tab to pull more.")
		return
	col1, col2 = st.columns(2)
	with col1:
		model_a = st.selectbox("Model A", models, index=0, key="cmp_a")
	with col2:
		model_b = st.selectbox("Model B", models, index=1 if len(models) > 1 else 0, key="cmp_b")
	text = st.text_area("Text to translate", height=160, placeholder="Enter text to compare translations…")
	detected = auto_detect_language(text) if text else None
	src_code, tgt_code = language_selectors(detected, key_ns="cmp_")
	if st.button("Run comparison"):
		if not text.strip():
			st.warning("Please enter text.")
			return
		if ollama is None:
			st.error("Ollama client not available.")
			return
		col_a, col_b = st.columns(2)
		with col_a:
			st.markdown(f"**{model_a}**")
			out_a = translate_text(
				text, model=model_a, src_lang=src_code, tgt_lang=tgt_code, temperature=temperature, stream=False
			)
			st.markdown(str(out_a))
		with col_b:
			st.markdown(f"**{model_b}**")
			out_b = translate_text(
				text, model=model_b, src_lang=src_code, tgt_lang=tgt_code, temperature=temperature, stream=False
			)
			st.markdown(str(out_b))


def main():
	st.set_page_config(page_title="Local LLM Translator (Ollama)", page_icon="🌐", layout="wide")
	st.title("Local LLM Translator 🌐")
	st.caption("Powered by your local Ollama models. No data leaves your machine.")

	if ollama is None:
		st.error(
			"Ollama Python client not found. Please run:\n\n"
			"pip install -r requirements.txt\n\n"
			"And ensure the Ollama service is running (`ollama serve`)."
		)
		return

	model, temperature, stream = page_sidebar()
	tab1, tab2, tab3, tab4 = st.tabs(["Translate Text", "Translate File", "Compare Models", "Models"])
	with tab1:
		ui_text_translation(model, temperature, stream)
	with tab2:
		ui_file_translation(model, temperature)
	with tab3:
		ui_compare_models(temperature)
	with tab4:
		ui_models_tab()

	st.divider()
	with st.expander("Notes and Tips"):
		st.markdown(
			"- If no models are listed, pull one first, e.g., `ollama pull qwen2.5:7b-instruct`.\n"
			"- PDF layout is not preserved; the output maintains text content and paragraphs.\n"
			"- For best results, choose instruction-tuned models (e.g., qwen2.5-instruct, mistral-instruct, llama3.x).\n"
			"- Set OLLAMA_HOST if your service is remote (e.g., `export OLLAMA_HOST=http://localhost:11434`)."
		)


if __name__ == "__main__":
	main()


